"""
AireLogix Credit Memo Generator
One-page institutional-grade Word document for lender distribution.
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD  = RGBColor(0xC8, 0xA9, 0x6E)
STEEL = RGBColor(0x8B, 0xA4, 0xBE)
CREAM = RGBColor(0xF0, 0xE8, 0xD8)

PRIORITY_LABELS = {
    "rate":         "Rate / Cost of Capital",
    "advance":      "Advance / LTV",
    "amortization": "Amortization Period",
    "prepayment":   "Prepayment Flexibility",
}


def _fmtd(v):
    if v is None or v == 0:
        return "—"
    v = float(v)
    if v >= 1_000_000:
        return f"${v/1_000_000:.2f}M"
    return f"${v:,.0f}"


def _fmtp(v, d=1):
    if not v and v != 0:
        return "—"
    return f"{float(v)*100:.{d}f}%"


def _fmtx(v):
    if not v and v != 0:
        return "—"
    return f"{float(v):.2f}x"


def _section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = GOLD
    run.font.name = "Calibri"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C8A96E")
    pBdr.append(bot)
    pPr.append(pBdr)


def _kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = Inches(3.2)
        vc.width = Inches(2.8)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(9)
        lr.font.name = "Calibri"
        lr.font.color.rgb = STEEL
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(9)
        vr.font.name = "Calibri"
        vr.bold = True
        vp.paragraph_format.space_before = Pt(2)
        vp.paragraph_format.space_after = Pt(2)
    return table


def generate_memo(analysis: dict, loan_prefs: dict = None,
                  borrower_profile: dict = None) -> bytes:
    from deals_engine import get_balloon

    lp  = loan_prefs or {}
    bp  = borrower_profile or {}
    a   = analysis
    t   = a.get("transaction", {})
    bs  = a.get("balanceSheet", {})
    inc = a.get("incomeNormalization", {})
    gd  = a.get("gdscr", {})
    ac  = a.get("aircraft", {})
    col = a.get("collateral", {})

    is_corp = a.get("dealType") in ("private_company", "corporate", "company")
    borrower_name  = a.get("borrowerName", "Identity Withheld")
    analysis_date  = a.get("analysisDate", "")
    deal_id        = a.get("applicationId", "")

    aircraft_str   = (ac.get("description") or
                      f"{ac.get('year','')} {ac.get('make','')} {ac.get('model','')}".strip())
    serial         = ac.get("serialNumber") or col.get("serialNumber") or "—"
    reg            = ac.get("registration") or col.get("registration") or "—"
    aftt           = ac.get("aftt") or col.get("aftt") or 0
    engine_program = ac.get("engineProgram") or col.get("engineProgram") or "—"
    fmv            = t.get("fmv") or col.get("fmv") or 0
    loan_amt       = t.get("loanAmount") or 0
    ac_year        = ac.get("year") or col.get("year") or 0
    ac_make        = ac.get("make") or col.get("make") or ""
    ac_model       = ac.get("model") or col.get("model") or ""

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Masthead ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("AIRELOGIX")
    r.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"; r.font.color.rgb = GOLD
    r2 = p.add_run("  ·  CONFIDENTIAL CREDIT PACKAGE")
    r2.font.size = Pt(8); r2.font.name = "Calibri"; r2.font.color.rgb = STEEL

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r3 = p2.add_run(f"{deal_id}  ·  {analysis_date}")
    r3.font.size = Pt(8); r3.font.name = "Calibri"; r3.font.color.rgb = STEEL

    # ── 1. Deal Overview ──────────────────────────────────────────────────────
    _section(doc, "1. Deal Overview")
    _kv_table(doc, [
        ("Borrower",              borrower_name),
        ("Borrower Type",         "Corporate" if is_corp else "Individual / HNWI"),
        ("Aircraft",              aircraft_str),
        ("Serial / Registration", f"{serial}  ·  {reg}"),
        ("AFTT",                  f"{int(aftt):,}" if aftt else "—"),
        ("Engine Program",        engine_program),
        ("Loan Amount",           _fmtd(loan_amt)),
        ("Analysis Date",         analysis_date),
    ])

    # ── 2. Transaction Summary & Borrower Preferences ─────────────────────────
    _section(doc, "2. Transaction Summary & Borrower Preferences")

    preferred_term = str(lp.get("preferredTerm") or t.get("termMonths") or "120")
    if preferred_term == "flexible":
        term_display = "Flexible / Open"
    elif preferred_term.isdigit():
        term_display = f"{round(int(preferred_term)/12)}yr ({preferred_term} mo)"
    else:
        term_display = f"{t.get('termMonths', 120)} mo"

    structure   = lp.get("structureType") or t.get("structure") or "Balloon"
    prepay_type = lp.get("prepayType") or ""
    notes_text  = lp.get("notes") or ""
    ltv_display = _fmtp(t.get("ltvVsFMV") or (loan_amt / fmv if fmv else t.get("ltv") or 0))

    tx_rows = [
        ("Loan Amount",          _fmtd(loan_amt)),
        ("Purchase Price",       _fmtd(t.get("purchasePrice", 0))),
        ("LTV vs FMV",           ltv_display),
        ("Term (Requested)",     term_display),
        ("Illustrative Rate",    _fmtp(t.get("illustrativeRate", 0), 2)),
        ("Monthly P&I",          _fmtd(t.get("monthlyPayment", 0))),
        ("Annual Aircraft D/S",  _fmtd(t.get("annualAircraftDS", 0))),
        ("Existing Annual D/S",  _fmtd(t.get("existingAnnualDS", 0))),
        ("Total Pro-Forma D/S",  _fmtd(t.get("totalProFormaDS", 0))),
        ("Structure",            structure),
    ]
    if prepay_type:
        prepay_label = "Open — no penalty" if prepay_type == "open" else prepay_type.replace("_", " ").title()
        tx_rows.append(("Prepayment",  prepay_label))

    _kv_table(doc, tx_rows)

    priority_ranking = lp.get("priorityRanking") or ["rate", "advance", "amortization", "prepayment"]
    if priority_ranking:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(8)
        ph.paragraph_format.space_after = Pt(3)
        rh = ph.add_run("Borrower Priorities — Ranked")
        rh.bold = True; rh.font.size = Pt(8.5); rh.font.name = "Calibri"; rh.font.color.rgb = GOLD
        for i, key in enumerate(priority_ranking):
            pi = doc.add_paragraph()
            pi.paragraph_format.space_before = Pt(1)
            pi.paragraph_format.space_after = Pt(1)
            pi.paragraph_format.left_indent = Inches(0.2)
            ri = pi.add_run(f"{i+1}.  {PRIORITY_LABELS.get(key, key.replace('_',' ').title())}")
            ri.font.size = Pt(9); ri.font.name = "Calibri"
            ri.bold = (i == 0)
            ri.font.color.rgb = CREAM if i == 0 else STEEL

    if notes_text:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(6)
        pn.paragraph_format.space_after = Pt(2)
        rn1 = pn.add_run("Borrower Notes:  ")
        rn1.bold = True; rn1.font.size = Pt(9); rn1.font.name = "Calibri"
        rn2 = pn.add_run(notes_text)
        rn2.italic = True; rn2.font.size = Pt(9); rn2.font.name = "Calibri"

    # ── 3. Borrower Background ────────────────────────────────────────────────
    _section(doc, "3. Borrower Background")
    narrative = bp.get("narrative") or ""
    if narrative:
        for chunk in narrative.split("\n"):
            chunk = chunk.strip()
            if not chunk:
                continue
            pn = doc.add_paragraph()
            pn.paragraph_format.space_before = Pt(2)
            pn.paragraph_format.space_after = Pt(4)
            rn = pn.add_run(chunk)
            rn.font.size = Pt(9); rn.font.name = "Calibri"
    else:
        pn = doc.add_paragraph()
        rn = pn.add_run("Borrower background research is in progress.")
        rn.italic = True; rn.font.size = Pt(9); rn.font.name = "Calibri"; rn.font.color.rgb = STEEL

    # ── 4. Detailed Financials ────────────────────────────────────────────────
    _section(doc, "4. Detailed Financials")
    tax_years = inc.get("taxYears") or []
    y1 = str(tax_years[0]) if tax_years else "Year 1"
    y2 = str(tax_years[1]) if len(tax_years) > 1 else "Year 2"

    if is_corp:
        fin_rows = [
            (f"Revenue — {y1}",              _fmtd(inc.get("revenueYear1", 0))),
            (f"Revenue — {y2}",              _fmtd(inc.get("revenueYear2", 0))),
            (f"EBITDA — {y1}",               _fmtd(inc.get("year1Total", 0))),
            (f"EBITDA — {y2}",               _fmtd(inc.get("year2Total", 0))),
            ("Qualifying EBITDA (Lower Yr)", _fmtd(inc.get("qualifyingIncome", 0))),
            ("EBITDA / Debt Service",        _fmtx(gd.get("gdscr", 0))),
            ("Financial Statement Quality",  (inc.get("financialStatementQuality") or "—").title()),
            ("",                             ""),
            ("Total Assets",                 _fmtd(bs.get("grossTotalAssets", 0))),
            ("Current Assets",               _fmtd(bs.get("totalCurrentAssets") or bs.get("tier1NetLiquid", 0))),
            ("Total Liabilities",            _fmtd(bs.get("totalLiabilities", 0))),
            ("Current Liabilities",          _fmtd(bs.get("totalCurrentLiabilities", 0))),
            ("Entity Net Worth",             _fmtd(bs.get("statedNetWorth", 0))),
            ("Equity / Loan Multiple",       _fmtx(bs.get("netWorthCoverage", 0))),
            ("Current Ratio",                _fmtx(bs.get("currentRatio") or bs.get("liquidityRatio", 0))),
            ("Leverage (Liabilities/Assets)",_fmtp(bs.get("leverageRatio", 0))),
            ("Debt / EBITDA (Pro Forma)",    _fmtx(bs.get("debtToEbitda", 0))),
        ]
        guarantors = a.get("guarantors", [])
        if guarantors:
            fin_rows.append(("", ""))
            fin_rows.append(("Personal Guarantors",
                             ", ".join(g.get("name", "—") for g in guarantors)))
    else:
        qualifying_year = inc.get("qualifyingYear") or "—"
        fin_rows = [
            (f"Qualifying Income — {y1}",     _fmtd(inc.get("year1Total", 0))),
            (f"Qualifying Income — {y2}",     _fmtd(inc.get("year2Total", 0))),
            ("Governing Year (Lower)",         str(qualifying_year)),
            ("Qualifying Income (Gross)",      _fmtd(inc.get("qualifyingIncome", 0))),
            ("After-Tax Qualifying Income",    _fmtd(inc.get("afterTaxQualifyingIncome", 0))),
            ("GDSCR (Pro-Forma)",              _fmtx(gd.get("gdscr", 0))),
            ("",                              ""),
            ("Net Liquid Assets",              _fmtd(bs.get("tier1NetLiquid") or bs.get("liquidAssets", 0))),
            ("Total Liabilities",             _fmtd(bs.get("totalLiabilities", 0))),
            ("Stated Net Worth",              _fmtd(bs.get("statedNetWorth", 0))),
            ("Net Worth / Loan Coverage",     _fmtx(bs.get("netWorthCoverage", 0))),
            ("Liquidity Ratio (Liquid/Loan)", _fmtx(bs.get("liquidityRatio", 0))),
        ]
        k1s = inc.get("k1Detail", [])
        if k1s:
            fin_rows += [("", ""),
                         ("K-1 Entities", f"{len(k1s)} entities")]
            for k in k1s:
                if not k.get("excluded"):
                    fin_rows.append((
                        f"  {k.get('entityName','—')} ({k.get('ownershipPct',0):.0f}%)",
                        f"Y1: {_fmtd(k.get('year1Box1',0))}  ·  Y2: {_fmtd(k.get('year2Box1',0))}"
                    ))
        debt = a.get("debtDetail") or []
        if debt:
            fin_rows.append(("", ""))
            for d in debt:
                fin_rows.append((
                    f"  {d.get('creditor','—')} ({d.get('accountType','—')})",
                    f"Balance: {_fmtd(d.get('balance',0))}  ·  Annual DS: {_fmtd(d.get('annualPayment',0))}"
                ))

    _kv_table(doc, fin_rows)

    # ── 5. Collateral & Balloon Projections ───────────────────────────────────
    _section(doc, "5. Collateral & Balloon Projections")
    _kv_table(doc, [
        ("Aircraft",        aircraft_str),
        ("Serial Number",   serial),
        ("Registration",    reg),
        ("AFTT",            f"{int(aftt):,}" if aftt else "—"),
        ("Engine Program",  engine_program),
        ("Current FMV",     _fmtd(fmv)),
        ("LTV vs FMV",      ltv_display),
    ])

    if fmv and loan_amt and ac_year:
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(8)
        ph.paragraph_format.space_after = Pt(4)
        rh = ph.add_run("Estimated Balloon Payments by Term")
        rh.bold = True; rh.font.size = Pt(8.5); rh.font.name = "Calibri"; rh.font.color.rgb = GOLD

        bt = doc.add_table(rows=1, cols=4)
        bt.style = "Table Grid"
        for i, hdr_txt in enumerate(["Term", "Est. FMV at Maturity", "Balloon @ 70% FMV", "Balloon / Loan"]):
            hc = bt.rows[0].cells[i]
            hr = hc.paragraphs[0].add_run(hdr_txt)
            hr.bold = True; hr.font.size = Pt(8.5); hr.font.name = "Calibri"

        for term_mo in [60, 84, 120, 180]:
            try:
                balloon   = get_balloon(fmv, loan_amt, int(ac_year), int(aftt),
                                        term_mo, ac_make, ac_model, engine_program)
                fmv_mat   = balloon / 0.70
                term_yrs  = term_mo // 12
                row_cells = bt.add_row().cells
                vals = [
                    f"{term_yrs}yr ({term_mo} mo)",
                    _fmtd(fmv_mat),
                    _fmtd(balloon),
                    _fmtp(balloon / loan_amt) if loan_amt else "—",
                ]
                for ci, val in enumerate(vals):
                    run = row_cells[ci].paragraphs[0].add_run(val)
                    run.font.size = Pt(9); run.font.name = "Calibri"
            except Exception:
                pass

    # ── Flags ─────────────────────────────────────────────────────────────────
    flags = a.get("flags", [])
    if flags:
        _section(doc, "Underwriting Flags")
        for f in flags:
            sev = f.get("severity", "INFO")
            pf = doc.add_paragraph()
            pf.paragraph_format.space_before = Pt(2)
            pf.paragraph_format.space_after = Pt(1)
            rf1 = pf.add_run(f"[{sev}]  ")
            rf1.bold = True; rf1.font.size = Pt(8.5); rf1.font.name = "Calibri"
            rf2 = pf.add_run(f.get("title", ""))
            rf2.bold = True; rf2.font.size = Pt(8.5); rf2.font.name = "Calibri"
            if f.get("detail"):
                pd = doc.add_paragraph()
                pd.paragraph_format.space_before = Pt(0)
                pd.paragraph_format.space_after = Pt(4)
                pd.paragraph_format.left_indent = Inches(0.2)
                rd = pd.add_run(f.get("detail", ""))
                rd.italic = True; rd.font.size = Pt(8.5); rd.font.name = "Calibri"

    # ── Footer ────────────────────────────────────────────────────────────────
    pf = doc.add_paragraph()
    pf.paragraph_format.space_before = Pt(16)
    rf = pf.add_run(
        "This document is confidential and intended solely for the credentialed lender institution "
        "to which it has been delivered. Prepared by AireLogix Aviation Finance Intelligence."
    )
    rf.italic = True; rf.font.size = Pt(7.5); rf.font.name = "Calibri"; rf.font.color.rgb = STEEL

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
